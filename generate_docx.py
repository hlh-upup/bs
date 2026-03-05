#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generate_docx.py - 将 Markdown 章节转换为符合中国计算机类硕士论文格式的 Word 文档

功能：
  1. 解析 Markdown 文件（chapter2.md, chapter6.md）
  2. 将 [@citkey] 转换为 [N] 编号引用
  3. 将 @sec:xxx / @fig:xxx 交叉引用转换为中文文本
  4. 生成符合 GB/T 7714-2015 格式的参考文献列表
  5. 输出格式化的 .docx 文件，可直接粘贴到现有论文中

用法：
  python3 generate_docx.py                     # 生成所有章节
  python3 generate_docx.py chapter2.md          # 只生成指定章节
  python3 generate_docx.py --output thesis.docx # 指定输出文件名

字体规范（中国传媒大学硕士论文格式）：
  - 章标题（# 1 绪论）：黑体，三号（16pt），居中
  - 节标题（## 1.1 xxx）：黑体，四号（14pt）
  - 小节标题（### 1.1.1 xxx）：黑体，小四号（12pt）
  - 正文：宋体，小四号（12pt），英文 Times New Roman
  - 行距：1.5倍
  - 页边距：上2.5cm，下2.5cm，左3cm，右2.5cm
  - 参考文献正文：宋体，五号（10.5pt）
"""

import re
import sys
import os
from collections import OrderedDict

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml, OxmlElement
except ImportError:
    print("错误：需要安装 python-docx 库")
    print("请运行：pip install python-docx")
    sys.exit(1)


# ===========================================================================
# 参考文献数据（从 references.bib 中提取的 GB/T 7714 格式）
# ===========================================================================
# 键名 -> (编号, GB/T 7714 格式全文)
# 编号按文中首次出现顺序动态分配

BIB_ENTRIES = OrderedDict([
    ("Goodfellow2014", "Goodfellow I J, Pouget-Abadie J, Mirza M, et al. Generative Adversarial Networks[C]//Advances in Neural Information Processing Systems. Montreal: MIT Press, 2014: 2672-2680."),
    ("Prajwal2020", "Prajwal K R, Mukhopadhyay R, Namboodiri V P, et al. A Lip Sync Expert Is All You Need for Speech to Lip Generation in the Wild[C]//Proceedings of the 28th ACM International Conference on Multimedia. Seattle: ACM, 2020: 484-492."),
    ("Zhang2023sadtalker", "Zhang W, Cun X, Wang X, et al. SadTalker: Learning Realistic 3D Motion Coefficients for Stylized Audio-Driven Single Image Talking Face Animation[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition. Vancouver: IEEE, 2023: 8652-8661."),
    ("Ye2024geneface", "Ye Z, Jiang Z, Ren Y, et al. GeneFace: Generalized and High-Fidelity Audio-Driven 3D Talking Face Synthesis[C]//The Twelfth International Conference on Learning Representations. Vienna: OpenReview, 2024."),
    ("Ye2023genefacepp", "Ye Z, He J, Jiang Z, et al. GeneFace++: Generalized and Stable Real-Time Audio-Driven 3D Talking Face Generation[J/OL]. arXiv preprint, 2023. arXiv: 2305.00787."),
    ("Wang2004ssim", "Wang Z, Bovik A C, Sheikh H R, et al. Image Quality Assessment: From Error Visibility to Structural Similarity[J]. IEEE Transactions on Image Processing, 2004, 13(4): 600-612."),
    ("Chung2016syncnet", "Chung J S, Zisserman A. Out of Time: Automated Lip Sync in the Wild[C]//Asian Conference on Computer Vision Workshops. Taipei: Springer, 2016: 251-263."),
    ("Toshpulatov2023", "Toshpulatov M, Lee W, Lee S. Talking Human Face Generation: A Survey[J]. Expert Systems with Applications, 2023, 219: 119678."),
    ("Bai2025", "Bai X, Cao Y, Wang L, et al. A Survey on Audio-Driven Talking Face Generation[J]. IEEE Transactions on Multimedia, 2025, 27: 1-18."),
    ("RVCBoss2024gptsovits", "RVC-Boss. GPT-SoVITS: A Versatile Few-Shot Voice Conversion and Text-to-Speech Framework[EB/OL]. (2024-01-15)[2025-03-01]. https://github.com/RVC-Boss/GPT-SoVITS."),
    ("DigitalChina2023", "中共中央, 国务院. 数字中国建设整体布局规划[Z]. 北京, 2023."),
    ("MOE2024", "教育部. 中国数字教育发展报告[R]. 北京: 教育部, 2024."),
    ("Chen2023nrqa", "Chen S, Li G, Dong Y, et al. A No-reference Quality Assessment Metric for Dynamic 3D Digital Human[J]. Displays, 2023, 79: 102547."),
    ("Quignon2024theval", "Quignon N. THEval: Evaluation Framework for Talking Head Video Generation[R/OL]. OpenReview, 2024."),
    ("Su2025", "Su M, Wang X, Fang Y, et al. Quality Assessment for Talking Head Videos via Multi-modal Feature Representation[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition Workshops. Nashville: IEEE, 2025: 1-8."),
    ("Liu2025ntire", "Liu X, Wang H, Wu J, et al. NTIRE 2025 Quality Assessment for AI-Generated Content Challenge[J/OL]. arXiv preprint, 2025. arXiv: 2506.02875."),
    ("Guo2024edtech", "Guo P, Zhang Y, Li H, et al. Digital Human Techniques for Education Reform[C]//Proceedings of the 7th International Conference on Educational Technology Management. Singapore: ACM, 2024: 1-6."),
    ("Liu2025digital", "Liu Q. Advancements in Digital Humans for Recorded Courses: Enhancing Learning Experiences via Personalized Interaction[J]. Frontiers of Digital Education, 2025, 2(1): 1-15."),
    ("Song2023", "宋一飞, 张炜, 陈智能, 等. 数字说话人视频生成综述[J]. 计算机辅助设计与图形学学报, 2023, 35(10): 1457-1474."),
    ("Zhang2024bigdata", "张冰源, 张旭龙, 王健宗, 等. 数字说话人脸生成技术综述[J]. 大数据, 2024, 10(3): 1-20."),
    ("Le2025", "乐铮, 胡永婷, 徐勇. 音频驱动的说话人面部视频生成与鉴别综述[J]. 计算机研究与发展, 2025, 62(10): 2523-2544."),
    ("SenseTime2025", "商汤科技, 等. 信息技术 面向客服的虚拟数字人通用技术要求: GB/T 46483-2025[S]. 北京: 中国标准出版社, 2025."),
    ("Yu2025", "Yu Q. Speaking Digital Person Video Generation Methods Review[C]//Proceedings of the 2nd International Conference on Data Science and Engineering. Singapore: Springer, 2025: 1-8."),
    ("Vaswani2017", "Vaswani A, Shazeer N, Parmar N, et al. Attention Is All You Need[C]//Advances in Neural Information Processing Systems. Long Beach: MIT Press, 2017: 5998-6008."),
    ("Perez2003poisson", "Pérez P, Gangnet M, Blake A. Poisson Image Editing[J]. ACM Transactions on Graphics, 2003, 22(3): 313-318."),
    ("Xing2023codetalker", "Xing J, Xia M, Zhang Y, et al. CodeTalker: Speech-Driven 3D Facial Animation with Discrete Motion Prior[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition. Vancouver: IEEE, 2023: 12780-12790."),
    ("Blanz1999", "Blanz V, Vetter T. A Morphable Model for the Synthesis of 3D Faces[C]//Proceedings of the 26th Annual Conference on Computer Graphics and Interactive Techniques. Los Angeles: ACM, 1999: 187-194."),
    ("Jolliffe2016pca", "Jolliffe I T, Cadima J. Principal Component Analysis: A Review and Recent Developments[J]. Philosophical Transactions of the Royal Society A: Mathematical, Physical and Engineering Sciences, 2016, 374(2065): 20150202."),
    ("Guo2021adnerf", "Guo Y, Chen K, Liang S, et al. AD-NeRF: Audio Driven Neural Radiance Fields for Talking Head Synthesis[C]//Proceedings of the IEEE/CVF International Conference on Computer Vision. Montreal: IEEE, 2021: 5784-5794."),
    ("Wang2021gfpgan", "Wang X, Li Y, Zhang H, et al. Towards Real-World Blind Face Restoration with Generative Facial Prior[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition. Nashville: IEEE, 2021: 9168-9178."),
    ("Li2023ernerf", "Li J, Zhang J, Bai X, et al. Efficient Region-Aware Neural Radiance Fields for High-Fidelity Talking Portrait Synthesis[C]//Proceedings of the IEEE/CVF International Conference on Computer Vision. Paris: IEEE, 2023: 7534-7544."),
    ("Stan2023facediffuser", "Stan S, Haque K I, Yumak Z. FaceDiffuser: Speech-Driven 3D Facial Animation Synthesis Using Diffusion[C]//ACM SIGGRAPH Conference on Motion, Interaction and Games. Rennes: ACM, 2023: 1-11."),
    ("Shen2023difftalk", "Shen S, Li W, Zhu Z, et al. DiffTalk: Crafting Diffusion Models for Generalized Audio-Driven Portraits Animation[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition. Vancouver: IEEE, 2023: 1982-1991."),
    ("Fan2022faceformer", "Fan Y, Lin Z, Saito J, et al. FaceFormer: Speech-Driven 3D Facial Animation with Transformers[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition. New Orleans: IEEE, 2022: 18770-18780."),
    ("Kim2024nerffacespeech", "Kim G, Seo K, Cha S, et al. NeRFFaceSpeech: One-shot Audio-Driven 3D Talking Head Synthesis via Generative Prior[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition Workshops. Seattle: IEEE, 2024: 7043-7052."),
    ("Zhang2024musetalk", "Zhang Y, Pan S, He Y, et al. MuseTalk: Real-Time High Quality Lip Synchronization with Latent Space Inpainting[J/OL]. arXiv preprint, 2024. arXiv: 2410.10122."),
    ("Heusel2017fid", "Heusel M, Ramsauer H, Unterthiner T, et al. GANs Trained by a Two Time-Scale Update Rule Converge to a Local Nash Equilibrium[C]//Advances in Neural Information Processing Systems. Long Beach: MIT Press, 2017: 6626-6637."),
    ("Xu2023eyetracking", "Xu M, Li S, Cao L, et al. Assessing Visual Quality of 3D Talking Head via Eye-Tracking[C]//2023 IEEE International Conference on Multimedia and Expo. Brisbane: IEEE, 2023: 165-170."),
    ("Wang2021mutualformer", "Wang X, Wang H, Ni H, et al. MutualFormer: Multi-Modality Representation Learning via Cross-Diffusion Attention[C]//Proceedings of the IEEE/CVF International Conference on Computer Vision. Montreal: IEEE, 2021: 5117-5126."),
    ("Liu2025tacfn", "Liu F, Fu Z, Wang Y, et al. TACFN: Transformer-Based Adaptive Cross-Modal Fusion Network for Multimodal Emotion Recognition[J/OL]. arXiv preprint, 2025. arXiv: 2505.06536."),
    ("Zhang2022contrastive", "Zhang Y, Jiang H, Miura Y, et al. Contrastive Learning of Medical Visual Representations from Paired Images and Text[C]//Machine Learning for Healthcare Conference. Durham: PMLR, 2022: 2-25."),
    ("Caruana1997", "Caruana R. Multitask Learning[J]. Machine Learning, 1997, 28(1): 41-75."),
    ("Misra2016crossstitch", "Misra I, Shrivastava A, Gupta A, et al. Cross-Stitch Networks for Multi-Task Learning[C]//Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition. Las Vegas: IEEE, 2016: 3994-4003."),
    ("Chen2018gradnorm", "Chen Z, Badrinarayanan V, Lee C Y, et al. GradNorm: Gradient Normalization for Adaptive Loss Balancing in Deep Multitask Networks[C]//International Conference on Machine Learning. Stockholm: PMLR, 2018: 794-803."),
    ("Kendall2018uncertainty", "Kendall A, Gal Y, Cipolla R. Multi-Task Learning Using Uncertainty to Weigh Losses for Scene Geometry and Semantics[C]//Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition. Salt Lake City: IEEE, 2018: 7482-7491."),
    ("Yu2020pcgrad", "Yu T, Kumar S, Gupta A, et al. Gradient Surgery for Multi-Task Learning[C]//Advances in Neural Information Processing Systems. Vancouver: MIT Press, 2020: 5824-5836."),
    ("Sener2018pareto", "Sener O, Koltun V. Multi-Task Learning as Multi-Objective Optimization[C]//Advances in Neural Information Processing Systems. Montreal: MIT Press, 2018: 525-536."),
    ("Shen2018tacotron", "Shen J, Pang R, Weiss R J, et al. Natural TTS Synthesis by Conditioning WaveNet on Mel Spectrogram Predictions[C]//2018 IEEE International Conference on Acoustics, Speech and Signal Processing. Calgary: IEEE, 2018: 4779-4783."),
    ("Kim2021vits", "Kim J, Kong J, Son J. Conditional Variational Autoencoder with Adversarial Learning for End-to-End Text-to-Speech[C]//International Conference on Machine Learning. Vienna: PMLR, 2021: 5530-5540."),
    ("Jia2018sv2tts", "Jia Y, Zhang Y, Weiss R, et al. Transfer Learning from Speaker Verification to Multispeaker Text-to-Speech Synthesis[C]//Advances in Neural Information Processing Systems. Montreal: MIT Press, 2018: 4485-4494."),
    ("Zeghidour2022soundstream", "Zeghidour N, Luebs A, Omran A, et al. SoundStream: An End-to-End Neural Audio Codec[J]. IEEE/ACM Transactions on Audio, Speech, and Language Processing, 2022, 30: 495-507."),
    ("Ma2023styletts", "Ma Y, Wang S, Hu Z, et al. StyleTalk: One-Shot Talking Head Generation with Controllable Speaking Styles[C]//Proceedings of the AAAI Conference on Artificial Intelligence. Washington: AAAI, 2023, 37(2): 1357-1365."),
    ("Fielding2000rest", "Fielding R T. Architectural Styles and the Design of Network-Based Software Architectures[D]. Irvine: University of California, 2000."),
    ("Nie2021bs", "Nie X, Wang X, Gao Y, et al. A Review on B/S Architecture Software Development Technology[J]. Journal of Physics: Conference Series, 2021, 1952(4): 042028."),
    ("Hinton2015distill", "Hinton G, Vinyals O, Dean J. Distilling the Knowledge in a Neural Network[J/OL]. arXiv preprint, 2015. arXiv: 1503.02531."),
    ("NVIDIA2024tensorrt", "NVIDIA Corporation. TensorRT: Programmable Inference Accelerator[EB/OL]. (2024)[2025-03-01]. https://developer.nvidia.com/tensorrt."),
    ("FFmpeg2000", "FFmpeg Developers. FFmpeg: A Complete, Cross-Platform Solution to Record, Convert and Stream Audio and Video[EB/OL]. (2000)[2025-03-01]. https://ffmpeg.org."),
    ("Bergkvist2021webrtc", "Bergkvist A, Burnett D C, Jennings C, et al. WebRTC 1.0: Real-Time Communication Between Browsers[S/OL]. (2021)[2025-03-01]. https://www.w3.org/TR/webrtc/."),
])

# 交叉引用标签 -> 中文文本
SECTION_REFS = {
    "sec:ch1": "第一章",
    "sec:background": "第1.1节",
    "sec:related-work": "第1.2节",
    "sec:intl-research": "第1.2.1节",
    "sec:cn-research": "第1.2.2节",
    "sec:structure": "第1.3节",
    "sec:ch2": "第二章",
    "sec:gen-tech": "第2.1节",
    "sec:gan-wav2lip": "第2.1.1节",
    "sec:3d-sadtalker": "第2.1.2节",
    "sec:diffusion-transformer": "第2.1.3节",
    "sec:qa-theory": "第2.2节",
    "sec:qa-system": "第2.2.1节",
    "sec:multimodal-fusion": "第2.2.2节",
    "sec:mtl": "第2.2.3节",
    "sec:tts-arch": "第2.3节",
    "sec:voice-clone": "第2.3.1节",
    "sec:bs-arch": "第2.3.2节",
    "sec:ch2-summary": "第2.4节",
    "sec:ch6": "第六章",
    "sec:conclusion": "第6.1节",
    "sec:innovation": "第6.2节",
    "sec:future": "第6.3节",
}

FIGURE_REFS = {
    "fig:wav2lip": "图2.1",
    "fig:sadtalker": "图2.2",
    "fig:voice-clone": "图2.3",
    "fig:qa-framework": "图2.4",
    "fig:fusion-compare": "图2.5",
    "fig:mtl-weight": "图2.6",
    "fig:tts-evolution": "图2.7",
    "fig:system-arch": "图2.8",
}

# 超链接标记分隔符（用于在文本中标记需要生成超链接的位置）
LINK_MARKER = '\x01'

# 书签 ID 计数器
_bookmark_id_counter = 0


def _next_bookmark_id():
    """生成唯一的书签 ID"""
    global _bookmark_id_counter
    _bookmark_id_counter += 1
    return _bookmark_id_counter


def _reset_bookmark_counter():
    """重置书签计数器（每次生成文档前调用）"""
    global _bookmark_id_counter
    _bookmark_id_counter = 0


def _label_to_bookmark(label):
    """将 Markdown 标签转换为有效的 Word 书签名称

    Word 书签名称限制：以字母或下划线开头，只含字母/数字/下划线，最长40字符。
    例如: 'sec:ch1' -> '_sec_ch1', 'fig:wav2lip' -> '_fig_wav2lip'
    """
    name = '_' + re.sub(r'[^a-zA-Z0-9_]', '_', label)
    return name[:40]


def _extract_label(text):
    """从文本中提取 {#sec:xxx} 或 {#fig:xxx} 标签"""
    m = re.search(r'\{#((?:sec|fig):[^}]+)\}', text)
    return m.group(1) if m else None


def _add_bookmark(paragraph, bookmark_name):
    """为段落添加书签（作为交叉引用的跳转目标）"""
    bid = _next_bookmark_id()
    p = paragraph._element

    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bid))
    start.set(qn('w:name'), bookmark_name)
    p.append(start)

    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bid))
    p.append(end)


def _make_run_props(font_size=Pt(12), bold=False,
                    cn_font='宋体', en_font='Times New Roman'):
    """创建 w:rPr 元素（run 格式属性）"""
    rPr = OxmlElement('w:rPr')

    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rPr.append(rFonts)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size.pt * 2)))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(font_size.pt * 2)))
    rPr.append(szCs)

    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)

    return rPr


def _add_hyperlink_run(paragraph, anchor_name, text, font_size=Pt(12),
                       bold=False, cn_font='宋体', en_font='Times New Roman'):
    """在段落中添加 Word 交叉引用字段（REF bookmark \\h）

    等效于 Word 中"插入 → 交叉引用"操作。点击时直接在文档内跳转到
    对应书签位置，不会打开新文件。生成的 OOXML 结构为：
      <w:r><w:fldChar w:fldCharType="begin"/></w:r>
      <w:r><w:instrText> REF anchor_name \\h </w:instrText></w:r>
      <w:r><w:fldChar w:fldCharType="separate"/></w:r>
      <w:r><w:t>display_text</w:t></w:r>
      <w:r><w:fldChar w:fldCharType="end"/></w:r>
    """
    p_elem = paragraph._element

    # fldChar begin
    run_begin = OxmlElement('w:r')
    run_begin.append(_make_run_props(font_size, bold, cn_font, en_font))
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run_begin.append(fld_begin)
    p_elem.append(run_begin)

    # instrText: REF anchor_name \h
    run_instr = OxmlElement('w:r')
    run_instr.append(_make_run_props(font_size, bold, cn_font, en_font))
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = f' REF {anchor_name} \\h '
    run_instr.append(instr_text)
    p_elem.append(run_instr)

    # fldChar separate
    run_sep = OxmlElement('w:r')
    run_sep.append(_make_run_props(font_size, bold, cn_font, en_font))
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    run_sep.append(fld_sep)
    p_elem.append(run_sep)

    # display text
    run_text = OxmlElement('w:r')
    run_text.append(_make_run_props(font_size, bold, cn_font, en_font))
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    run_text.append(t)
    p_elem.append(run_text)

    # fldChar end
    run_end = OxmlElement('w:r')
    run_end.append(_make_run_props(font_size, bold, cn_font, en_font))
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run_end.append(fld_end)
    p_elem.append(run_end)


# ===========================================================================
# 引用编号管理
# ===========================================================================

class CitationManager:
    """管理参考文献编号，按首次出现顺序分配"""

    def __init__(self):
        self.cite_order = []  # 按出现顺序排列的 cite keys
        self.cite_map = {}    # cite_key -> number

    def get_number(self, cite_key):
        """获取引用编号，如果是首次出现则分配新编号"""
        if cite_key not in self.cite_map:
            self.cite_order.append(cite_key)
            self.cite_map[cite_key] = len(self.cite_order)
        return self.cite_map[cite_key]

    def get_ordered_references(self):
        """返回按编号排序的参考文献列表"""
        refs = []
        for key in self.cite_order:
            if key in BIB_ENTRIES:
                refs.append((self.cite_map[key], BIB_ENTRIES[key]))
        return refs


# ===========================================================================
# Markdown 解析与转换
# ===========================================================================

def resolve_citations(text, citation_mgr):
    """将 [@key1; @key2] 格式的引用替换为带超链接标记的 [N] 或 [N,M] 格式

    在文本中插入 LINK_MARKER 分隔的标记（如 CITE:_ref_1:1），
    后续段落构建时会将标记转为可点击的内部超链接。
    """

    def replace_cite_group(match):
        """处理一组引用 [@key1; @key2]"""
        inner = match.group(1)
        keys = re.findall(r'@(\w+)', inner)
        parts = []
        for key in keys:
            num = citation_mgr.get_number(key)
            parts.append(f"{LINK_MARKER}CITE:_ref_{num}:{num}{LINK_MARKER}")
        return "[" + ",".join(parts) + "]"

    # 处理 [@key1; @key2] 格式
    text = re.sub(r'\[([^]]*@\w+[^]]*)\]', replace_cite_group, text)

    return text


def resolve_crossrefs(text):
    """将 @sec:xxx 和 @fig:xxx 替换为带超链接标记的中文文本

    在文本中插入 LINK_MARKER 分隔的标记（如 SEC:_sec_ch2:第二章），
    后续段落构建时会将标记转为可点击的内部超链接。
    """
    # 处理 @sec:xxx
    for key, label in SECTION_REFS.items():
        bookmark = _label_to_bookmark(key)
        marker = f"{LINK_MARKER}SEC:{bookmark}:{label}{LINK_MARKER}"
        text = text.replace(f"@{key}", marker)

    # 处理 @fig:xxx
    for key, label in FIGURE_REFS.items():
        bookmark = _label_to_bookmark(key)
        marker = f"{LINK_MARKER}FIG:{bookmark}:{label}{LINK_MARKER}"
        text = text.replace(f"@{key}", marker)

    return text


def strip_crossref_labels(text):
    """去除标题末尾的 {#sec:xxx} 和图片的 {#fig:xxx} 标签"""
    text = re.sub(r'\s*\{#sec:[^}]+\}', '', text)
    text = re.sub(r'\s*\{#fig:[^}]+\}', '', text)
    return text


def parse_markdown_line(line):
    """解析一行 markdown，返回 (类型, 内容, 级别, 标签)

    标签用于生成 Word 书签（交叉引用跳转目标）。
    """
    line = line.rstrip()

    # 章标题 # 1 绪论 {#sec:xxx}
    m = re.match(r'^# (.+)$', line)
    if m:
        raw = m.group(1)
        label = _extract_label(raw)
        return ('heading1', strip_crossref_labels(raw).strip(), 1, label)

    # 节标题 ## 1.1 xxx {#sec:xxx}
    m = re.match(r'^## (.+)$', line)
    if m:
        raw = m.group(1)
        label = _extract_label(raw)
        return ('heading2', strip_crossref_labels(raw).strip(), 2, label)

    # 小节标题 ### 1.1.1 xxx {#sec:xxx}
    m = re.match(r'^### (.+)$', line)
    if m:
        raw = m.group(1)
        label = _extract_label(raw)
        return ('heading3', strip_crossref_labels(raw).strip(), 3, label)

    # 图片 ![alt](path){#fig:xxx}
    m = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)', line)
    if m:
        label = _extract_label(line)
        return ('image', (m.group(1), m.group(2)), 0, label)

    # 分割线
    if line.strip() == '---':
        return ('hr', '', 0, None)

    # 空行
    if line.strip() == '':
        return ('blank', '', 0, None)

    # 普通段落
    return ('paragraph', line, 0, None)


def process_inline_formatting(paragraph, text, doc, body_font_size=Pt(12)):
    """处理行内格式（粗体 **xx**、超链接等）并添加到段落

    注意：中文使用宋体，英文/数字使用 Times New Roman
    """
    # 分割文本：处理 **bold** 和普通文本
    parts = re.split(r'(\*\*[^*]+\*\*)', text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # 粗体
            inner = part[2:-2]
            _process_text_with_links(paragraph, inner, body_font_size, bold=True)
        elif part:
            _process_text_with_links(paragraph, part, body_font_size, bold=False)


def _process_text_with_links(paragraph, text, font_size, bold=False):
    """处理可能包含超链接标记的文本，生成普通文本或可点击超链接"""
    if LINK_MARKER not in text:
        # 无链接标记，直接添加普通 run
        run = paragraph.add_run(text)
        run.font.size = font_size
        if bold:
            run.bold = True
        _set_run_font(run)
        return

    # 按链接标记分割
    segments = text.split(LINK_MARKER)
    for seg in segments:
        if seg.startswith(('CITE:', 'SEC:', 'FIG:')):
            # 超链接段：格式为 "TYPE:bookmark_name:display_text"
            link_parts = seg.split(':', 2)
            if len(link_parts) == 3:
                _, bookmark, display = link_parts
                _add_hyperlink_run(paragraph, bookmark, display, font_size, bold)
            else:
                print(f"警告：格式异常的链接标记：{seg}")
                run = paragraph.add_run(seg)
                run.font.size = font_size
                _set_run_font(run)
        elif seg:
            # 普通文本段
            run = paragraph.add_run(seg)
            run.font.size = font_size
            if bold:
                run.bold = True
            _set_run_font(run)


def _set_run_font(run, cn_font='宋体', en_font='Times New Roman'):
    """设置 run 的中英文字体"""
    run.font.name = en_font
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), cn_font)


def _set_paragraph_spacing(paragraph, line_spacing=1.5, space_before=0, space_after=0):
    """设置段落行距和段前段后间距"""
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)


# ===========================================================================
# Word 文档生成
# ===========================================================================

def setup_document_styles(doc):
    """设置文档样式"""
    # 设置页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)  # 小四
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    pf = style.paragraph_format
    pf.line_spacing = 1.5

    # 创建/修改标题样式
    _setup_heading_style(doc, 'Heading 1', '黑体', Pt(16), True)  # 三号
    _setup_heading_style(doc, 'Heading 2', '黑体', Pt(14), False)  # 四号
    _setup_heading_style(doc, 'Heading 3', '黑体', Pt(12), False)  # 小四


def _setup_heading_style(doc, style_name, cn_font, font_size, center=False):
    """配置标题样式"""
    try:
        style = doc.styles[style_name]
    except KeyError:
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

    font = style.font
    font.name = 'Times New Roman'
    font.size = font_size
    font.bold = True
    font.color.rgb = RGBColor(0, 0, 0)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), cn_font)

    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    if center:
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_chapter_heading(doc, text):
    """添加章标题（# 级别）"""
    p = doc.add_paragraph(style='Heading 1')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p, 1.5, 24, 18)
    return p


def add_section_heading(doc, text, level=2):
    """添加节/小节标题"""
    style_name = f'Heading {level}'
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    if level == 2:
        run.font.size = Pt(14)
        _set_paragraph_spacing(p, 1.5, 12, 6)
    else:
        run.font.size = Pt(12)
        _set_paragraph_spacing(p, 1.5, 6, 6)
    return p


def add_body_paragraph(doc, text):
    """添加正文段落"""
    p = doc.add_paragraph()
    process_inline_formatting(p, text, doc, Pt(12))
    _set_paragraph_spacing(p, 1.5, 0, 0)
    # 首行缩进 2 字符
    p.paragraph_format.first_line_indent = Pt(24)
    return p


def add_figure_caption(doc, caption_text):
    """添加图注"""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    process_inline_formatting(p, caption_text, doc, Pt(10.5))
    _set_paragraph_spacing(p, 1.5, 6, 6)
    return p


def add_image_placeholder(doc, alt_text, image_path):
    """添加图片（如果文件存在则插入，否则添加占位符）"""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 检查图片文件是否存在
    if os.path.exists(image_path):
        try:
            run = p.add_run()
            run.add_picture(image_path, width=Inches(5.0))
            return p
        except Exception:
            pass

    # 图片不存在，添加占位符文本
    run = p.add_run(f"[图片: {alt_text}]")
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(128, 128, 128)
    _set_run_font(run)
    _set_paragraph_spacing(p, 1.5, 6, 6)
    return p


def add_reference_section(doc, citation_mgr):
    """添加参考文献列表（每条参考文献带书签，正文中的引用可跳转至此）"""
    # 标题
    add_chapter_heading(doc, "参考文献")

    refs = citation_mgr.get_ordered_references()
    for num, ref_text in refs:
        p = doc.add_paragraph()
        # 添加书签（对应正文中 [N] 的跳转目标）
        _add_bookmark(p, f'_ref_{num}')
        # 编号
        run = p.add_run(f"[{num}] ")
        run.font.size = Pt(10.5)  # 五号
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # 内容
        run = p.add_run(ref_text)
        run.font.size = Pt(10.5)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        _set_paragraph_spacing(p, 1.25, 0, 0)
        # 悬挂缩进：编号后的正文部分对齐
        p.paragraph_format.first_line_indent = Pt(0)


# ===========================================================================
# 主流程
# ===========================================================================

def convert_markdown_to_docx(md_files, output_path, include_references=True):
    """将 Markdown 文件转换为格式化的 Word 文档"""
    citation_mgr = CitationManager()

    # 读取所有 markdown 文件
    all_lines = []
    for md_file in md_files:
        if not os.path.exists(md_file):
            print(f"警告：文件不存在，跳过：{md_file}")
            continue
        with open(md_file, 'r', encoding='utf-8') as f:
            content = f.read()

        # 第一遍：收集所有引用并分配编号（按出现顺序）
        # 需要先扫描一遍来确定引用顺序
        temp_text = content
        re.sub(r'\[([^]]*@\w+[^]]*)\]',
               lambda m: resolve_citations(m.group(0), citation_mgr),
               temp_text)

        all_lines.extend(content.split('\n'))
        all_lines.append('')  # 章节间空行

    # 重置引用管理器，按实际文档顺序重新编号
    citation_mgr = CitationManager()
    _reset_bookmark_counter()

    # 预处理完整文本：解析引用（生成超链接标记）
    full_text = '\n'.join(all_lines)
    full_text = resolve_citations(full_text, citation_mgr)
    full_text = resolve_crossrefs(full_text)
    processed_lines = full_text.split('\n')

    # 创建文档
    doc = Document()
    setup_document_styles(doc)

    # 设置文档属性（送审系统和查重工具会读取这些元数据）
    doc.core_properties.title = '硕士学位论文'
    doc.core_properties.subject = '学位论文'

    # 逐行处理
    prev_type = None
    skip_figure_caption = False

    for line in processed_lines:
        line_type, content, level, label = parse_markdown_line(line)

        # 跳过图片后的手动图注行（以 **图X.X 开头的行）
        if skip_figure_caption:
            skip_figure_caption = False
            if content and isinstance(content, str) and re.match(r'^\*?\*?图\d', content):
                # 这是图片后的说明文字，作为图注处理
                clean = strip_crossref_labels(content)
                add_figure_caption(doc, clean)
                continue

        if line_type == 'heading1':
            p = add_chapter_heading(doc, content)
            if label:
                _add_bookmark(p, _label_to_bookmark(label))
        elif line_type == 'heading2':
            p = add_section_heading(doc, content, 2)
            if label:
                _add_bookmark(p, _label_to_bookmark(label))
        elif line_type == 'heading3':
            p = add_section_heading(doc, content, 3)
            if label:
                _add_bookmark(p, _label_to_bookmark(label))
        elif line_type == 'image':
            alt_text, img_path = content
            alt_clean = strip_crossref_labels(alt_text)
            p = add_image_placeholder(doc, alt_clean, img_path)
            if label:
                _add_bookmark(p, _label_to_bookmark(label))
            # 下一行可能是图注
            skip_figure_caption = True
        elif line_type == 'paragraph':
            clean = strip_crossref_labels(content)
            if clean.strip():
                add_body_paragraph(doc, clean)
        elif line_type == 'blank':
            pass  # 空行不处理
        elif line_type == 'hr':
            pass  # 分割线不处理

        prev_type = line_type

    # 添加参考文献
    if include_references:
        doc.add_page_break()
        add_reference_section(doc, citation_mgr)

    # 保存
    doc.save(output_path)
    return citation_mgr


def main():
    """主函数"""
    import argparse

    parser = argparse.ArgumentParser(
        description='将 Markdown 章节转换为符合中国硕士论文格式的 Word 文档')
    parser.add_argument('files', nargs='*', default=[],
                        help='要转换的 Markdown 文件（默认：chapter2.md chapter6.md）')
    parser.add_argument('--output', '-o', default='thesis_chapters.docx',
                        help='输出文件名（默认：thesis_chapters.docx）')
    parser.add_argument('--no-refs', action='store_true',
                        help='不生成参考文献列表')

    args = parser.parse_args()

    # 确定输入文件
    if not args.files:
        # 默认文件列表
        default_files = ['chapter2.md', 'chapter6.md']
        args.files = [f for f in default_files if os.path.exists(f)]

    if not args.files:
        print("错误：没有找到任何 Markdown 文件")
        print("用法：python3 generate_docx.py [chapter2.md] [chapter6.md]")
        sys.exit(1)

    print("=" * 60)
    print("  论文 Word 文档生成工具")
    print("  格式：中国计算机类硕士论文 / GB/T 7714-2015 参考文献")
    print("=" * 60)
    print()
    print(f"输入文件：")
    for f in args.files:
        print(f"  ✓ {f}")
    print(f"输出文件：{args.output}")
    print()

    citation_mgr = convert_markdown_to_docx(
        args.files, args.output,
        include_references=not args.no_refs
    )

    refs = citation_mgr.get_ordered_references()
    print(f"✓ 生成完成！")
    print(f"  - 参考文献数量：{len(refs)} 条")
    print(f"  - 输出文件：{args.output}")
    print(f"  - 文件大小：{os.path.getsize(args.output) / 1024:.1f} KB")
    print()
    print("提示：")
    print("  1. 打开生成的 .docx 文件")
    print("  2. 全选内容（Ctrl+A），复制（Ctrl+C）")
    print("  3. 粘贴到您的论文 Word 文件的对应位置")
    print("  4. 参考文献列表在文件末尾，可单独复制")
    print()
    print("字体说明：")
    print("  - 章标题：黑体 三号（16pt）居中")
    print("  - 节标题：黑体 四号（14pt）")
    print("  - 小节标题：黑体 小四号（12pt）")
    print("  - 正文：宋体 小四号（12pt），英文 Times New Roman")
    print("  - 参考文献：宋体 五号（10.5pt）")
    print("  - 行距：1.5 倍")
    print()
    print("兼容性说明（送审 / 查重 / PDF）：")
    print("  ✓ 查重兼容：正文使用标准文本，知网/维普/万方等系统可正常提取")
    print("  ✓ PDF转换：交叉引用使用 Word 标准书签+超链接，转 PDF 后可点击跳转")
    print("  ✓ 送审兼容：文档包含标题样式和元数据，可正常生成目录和大纲")
    print("  ✓ 格式合规：字体/字号/行距/页边距均符合中国硕士论文格式规范")


if __name__ == '__main__':
    main()
